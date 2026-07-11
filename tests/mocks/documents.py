"""Generators for document test fixtures.

Every fixture is built in-process from a library, so the suite carries no binary
blobs and the bytes are reproducible. reportlab writes PDFs, python-docx writes
DOCX; the text formats are just strings.
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from collections.abc import Sequence


def make_pdf(
    pages: Sequence[Sequence[tuple[str, int]]],
    *,
    encrypt: str | None = None,
) -> bytes:
    """Build a PDF. Each page is a sequence of ``(text, font_size)`` lines.

    A larger font size makes a line render as a heading in the extractor, so a test
    can control what the heading detection sees. ``encrypt`` sets an owner password.
    """
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    if encrypt is not None:
        pdf.setEncrypt(_encrypt(encrypt))
    for page in pages:
        # Spread the lines across the full page height so the first sits in the top
        # margin and the last in the bottom margin — the geometry a real header and
        # footer have, which is what the chrome detector keys on.
        top, bottom = 770.0, 60.0
        step = (top - bottom) / max(len(page) - 1, 1)
        for index, (text, size) in enumerate(page):
            pdf.setFont("Helvetica-Bold" if size >= 16 else "Helvetica", size)
            pdf.drawString(72, top - index * step, text)
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _encrypt(password: str) -> object:
    from reportlab.lib import pdfencrypt

    return pdfencrypt.StandardEncryption(password, canPrint=0)


def make_docx(
    blocks: Sequence[tuple[str, str]], *, table: Sequence[Sequence[str]] | None = None
) -> bytes:
    """Build a DOCX. Each block is ``(style, text)`` — e.g. ``("Heading 1", "Title")``."""
    import docx

    document = docx.Document()
    for style, text in blocks:
        if style.lower().startswith("heading"):
            document.add_heading(text, level=int(style.split()[-1]))
        elif style == "title":
            document.add_heading(text, level=0)
        elif style.lower().startswith("list"):
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)
    if table:
        rows, cols = len(table), len(table[0])
        docx_table = document.add_table(rows=rows, cols=cols)
        for r, row in enumerate(table):
            for c, cell in enumerate(row):
                docx_table.rows[r].cells[c].text = cell

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_with_hyperlink(text: str, url: str) -> bytes:
    """A DOCX whose one paragraph contains a real hyperlink relationship."""
    import docx
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    document = docx.Document()
    paragraph = document.add_paragraph("See ")
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_text = OxmlElement("w:t")
    run_text.text = text
    run.append(run_text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def to_data_uri(data: bytes, media_type: str) -> str:
    import base64

    return f"data:{media_type};base64,{base64.b64encode(data).decode()}"


def to_base64(data: bytes) -> str:
    import base64

    return base64.b64encode(data).decode()
