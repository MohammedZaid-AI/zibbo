"""Reproducible document-extraction benchmarks.

    python -m benchmarks.documents

Every document is generated in-process from a fixed seed, so the byte and token
figures are identical on every machine. Runtime is the median of a few runs, and
peak memory is measured with tracemalloc — the allocation attributed to the
extraction, not whatever RSS the allocator happens to hold.

The number that matters most is **tokens after / tokens before**. A base64 PDF is
not just large; it tokenizes atrociously, because the tokenizer has no idea it is
looking at a document. Extraction replaces it with prose the model can actually read.
"""

from __future__ import annotations

import base64
import io
import random
import statistics
import time
import tracemalloc
from dataclasses import dataclass

from gateway.documents import build_document_registry
from gateway.documents.models import DocumentFormat
from gateway.tokenizers import TokenCounterFactory

SEED = 20260711
REPEAT = 5


@dataclass(frozen=True, slots=True)
class Row:
    name: str
    fmt: str
    input_bytes: int
    base64_tokens: int
    output_tokens: int
    token_reduction_pct: float
    median_ms: float
    peak_mib: float


_WORDS = [
    "revenue",
    "margin",
    "quarter",
    "growth",
    "region",
    "product",
    "customer",
    "segment",
    "forecast",
    "pipeline",
    "operating",
    "income",
    "expense",
    "strategy",
    "market",
    "share",
    "adoption",
    "retention",
    "expansion",
]


def _rng() -> random.Random:
    return random.Random(SEED)


def _sentence(rng: random.Random, n: int = 12) -> str:
    return " ".join(rng.choice(_WORDS) for _ in range(n)).capitalize() + "."


def _make_pdf(pages: int) -> bytes:
    from reportlab.pdfgen import canvas

    rng = _rng()
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    for page in range(pages):
        pdf.setFont("Helvetica-Bold", 18)
        pdf.drawString(72, 760, f"Section {page + 1}")
        pdf.setFont("Helvetica", 11)
        y = 730
        for _ in range(30):
            pdf.drawString(72, y, _sentence(rng))
            y -= 16
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _make_docx(sections: int) -> bytes:
    import docx

    rng = _rng()
    document = docx.Document()
    for section in range(sections):
        document.add_heading(f"Section {section + 1}", level=1)
        for _ in range(6):
            document.add_paragraph(_sentence(rng, 18))
        table = document.add_table(rows=4, cols=3)
        for row in table.rows:
            for cell in row.cells:
                cell.text = rng.choice(_WORDS)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_csv(rows: int) -> bytes:
    rng = _rng()
    lines = ["id,region,product,revenue,units,margin"]
    for index in range(rows):
        lines.append(
            f"{index},{rng.choice(_WORDS)},{rng.choice(_WORDS)},"
            f"{rng.randint(1000, 99999)},{rng.randint(1, 500)},{round(rng.random(), 3)}"
        )
    return "\n".join(lines).encode()


def _make_xml(records: int) -> bytes:
    rng = _rng()
    items = "".join(
        f'<record id="{i}"><region>{rng.choice(_WORDS)}</region>'
        f"<revenue>{rng.randint(1000, 99999)}</revenue>"
        f"<note>{_sentence(rng, 8)}</note></record>"
        for i in range(records)
    )
    return f"<?xml version='1.0'?><report>{items}</report>".encode()


def _benchmark(name: str, data: bytes, fmt: DocumentFormat, counter: object) -> Row:
    registry = build_document_registry()
    extractor = registry.for_format(fmt)
    assert extractor is not None

    result = extractor.extract(data, fmt)  # warm
    markdown = result.markdown or ""

    timings: list[float] = []
    for _ in range(REPEAT):
        started = time.perf_counter()
        extractor.extract(data, fmt)
        timings.append((time.perf_counter() - started) * 1000)

    tracemalloc.start()
    extractor.extract(data, fmt)
    _, peak = tracemalloc.get_traced_memory()
    tracemalloc.stop()

    base64_text = base64.b64encode(data).decode()
    base64_tokens = counter.count(base64_text)  # type: ignore[attr-defined]
    output_tokens = counter.count(markdown)  # type: ignore[attr-defined]

    return Row(
        name=name,
        fmt=fmt.value,
        input_bytes=len(data),
        base64_tokens=base64_tokens,
        output_tokens=output_tokens,
        token_reduction_pct=(
            round((base64_tokens - output_tokens) / base64_tokens * 100, 1)
            if base64_tokens
            else 0.0
        ),
        median_ms=round(statistics.median(timings), 1),
        peak_mib=round(peak / 1_048_576, 1),
    )


def main() -> None:
    counter = TokenCounterFactory().for_model("gpt-4o-mini")
    print(f"token counter: {counter.name} (exact={counter.exact})\n")

    rows = [
        _benchmark("10-page PDF", _make_pdf(10), DocumentFormat.PDF, counter),
        _benchmark("100-page PDF", _make_pdf(100), DocumentFormat.PDF, counter),
        _benchmark("DOCX report", _make_docx(20), DocumentFormat.DOCX, counter),
        _benchmark("CSV 5k rows", _make_csv(5000), DocumentFormat.CSV, counter),
        _benchmark("XML 2k records", _make_xml(2000), DocumentFormat.XML, counter),
    ]

    header = (
        f"{'document':<18} {'fmt':<5} {'in KB':>8} {'b64 tok':>9} "
        f"{'out tok':>9} {'reduced':>8} {'median':>9} {'peak MiB':>9}"
    )
    print(header)
    print("-" * len(header))
    for row in rows:
        print(
            f"{row.name:<18} {row.fmt:<5} {row.input_bytes / 1024:>8.1f} "
            f"{row.base64_tokens:>9,} {row.output_tokens:>9,} {row.token_reduction_pct:>7.1f}% "
            f"{row.median_ms:>8.1f}ms {row.peak_mib:>9.1f}"
        )


if __name__ == "__main__":
    main()
