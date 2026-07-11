"""DOCX -> Markdown, via python-docx.

A DOCX carries its structure explicitly — a heading *is* labelled ``Heading 1``, a
list item *is* labelled ``List Bullet`` — so extraction is reading the labels rather
than inferring from geometry as with a PDF. That makes it the most faithful of the
binary formats: headings, paragraphs, lists, tables and hyperlinks all survive.

The one subtlety is order. ``document.paragraphs`` and ``document.tables`` are
separate sequences, so iterating them independently would tear tables out of the
prose that surrounds them. The body's XML child order is walked instead, so a table
appears exactly where it sits in the document.
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING, ClassVar, Final

from gateway.documents.base import DocumentExtractor
from gateway.documents.models import DocumentFormat

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph

_W_NS: Final = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_MAX_HEADING: Final = 6


class DocxExtractor(DocumentExtractor):
    name: ClassVar[str] = "docx"
    formats: ClassVar[frozenset[DocumentFormat]] = frozenset({DocumentFormat.DOCX})

    def _extract(self, data: bytes, fmt: DocumentFormat) -> str | None:
        import docx

        document = docx.Document(io.BytesIO(data))
        blocks = [block for block in _iter_blocks(document) if block.strip()]
        return "\n\n".join(blocks) or None


def _iter_blocks(document: DocxDocument) -> list[str]:
    """Walk the body in document order, rendering paragraphs and tables in place."""
    blocks: list[str] = []
    parent = document.element.body
    tables = iter(document.tables)
    paragraphs = iter(document.paragraphs)

    for child in parent.iterchildren():
        if child.tag == f"{_W_NS}p":
            rendered = _paragraph(next(paragraphs, None))
            if rendered:
                blocks.append(rendered)
        elif child.tag == f"{_W_NS}tbl":
            rendered = _table(next(tables, None))
            if rendered:
                blocks.append(rendered)
    return blocks


def _paragraph(paragraph: Paragraph | None) -> str:
    if paragraph is None:
        return ""
    text = _runs_with_links(paragraph)
    if not text.strip():
        return ""

    style = (paragraph.style.name if paragraph.style else "") or ""
    lowered = style.lower()

    if lowered.startswith("heading"):
        level = _heading_level(style)
        return f"{'#' * level} {text}"
    if lowered.startswith("title"):
        return f"# {text}"
    if "list bullet" in lowered:
        return f"- {text}"
    if "list number" in lowered:
        return f"1. {text}"
    if lowered.startswith("quote") or "intense quote" in lowered:
        return f"> {text}"
    return text


def _heading_level(style: str) -> int:
    digits = "".join(ch for ch in style if ch.isdigit())
    if not digits:
        return 2
    return max(1, min(_MAX_HEADING, int(digits)))


def _runs_with_links(paragraph: Paragraph) -> str:
    """Reconstruct the paragraph text, turning hyperlinks into Markdown links.

    python-docx exposes runs but folds hyperlink runs oddly, so the paragraph XML is
    walked to pair link text with its relationship target.
    """
    parts: list[str] = []
    part = paragraph._p
    rels = paragraph.part.rels

    for node in part.iterchildren():
        tag = node.tag
        if tag == f"{_W_NS}r":
            parts.append(_run_text(node))
        elif tag == f"{_W_NS}hyperlink":
            text = "".join(_run_text(run) for run in node.iterchildren(f"{_W_NS}r"))
            rel_id = node.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
            )
            target = rels[rel_id].target_ref if rel_id and rel_id in rels else ""
            parts.append(f"[{text}]({target})" if target and text else text)
    return "".join(parts).strip()


def _run_text(run_element: object) -> str:
    return "".join(node.text or "" for node in run_element.iterchildren(f"{_W_NS}t"))  # type: ignore[attr-defined]


def _table(table: Table | None) -> str:
    if table is None:
        return ""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("|", r"\|").replace("\n", " ") for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""

    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header, *body = padded
    lines = ["| " + " | ".join(header) + " |", "|" + "---|" * width]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)
